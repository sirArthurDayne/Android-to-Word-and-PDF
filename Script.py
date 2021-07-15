import os
import docx
import sys
import docx2pdf

# ---------- Declaraciones ----------
user = input("Ingresa el user de la PC: ")
operation = input("\n1. android to word\n 2. word to pdf\n")
ruta_de_proyectos = "C:\\Users\\{}\\AndroidStudioProjects".format(user)
escritorio = "C:\\Users\\{}\\Desktop".format(user)
try:
    os.mkdir(os.path.join(escritorio, "Android a Word"))
except:
    pass

def filtro_por_extension(nombre):
    extensiones_a_incluir = [".java", ".xml", ".png", ".jpg", ".mp3", ".mp4"]
    nombre_extension = os.path.splitext(nombre)[-1].lower()
    return nombre_extension in extensiones_a_incluir or nombre == "build.gradle"

def filtro_por_archivo(nombre):
    excluir = ["ic_launcher_background.xml", "ic_launcher_foreground.xml", "ExampleUnitTest",
               "ExampleInstrumentedTest.java", "BuildConfig.java", "ExampleUnitTest.java"]
    return nombre not in excluir

def filtro_por_carpeta(root, proyecto_i):
    return (root.endswith("\layout") or root.endswith("\main")
            or root.endswith(proyecto_i) or root.endswith("drawable")
            or root.endswith("drawable-v24") or root.endswith("anim")
            or root.endswith("raw") or "\main\java" in root)

def es_imagen(nombre):
    extensiones = [".png", ".jpg"]
    nombre_extension = os.path.splitext(nombre)[-1].lower()
    return nombre_extension in extensiones

def es_mp3(nombre):
    extensiones = [".mp3"]
    nombre_extension = os.path.splitext(nombre)[-1].lower()
    return nombre_extension in extensiones

def es_mp4(nombre):
    extensiones = [".mp4"]
    nombre_extension = os.path.splitext(nombre)[-1].lower()
    return nombre_extension in extensiones


def es_docx(nombre):
    return nombre[-5:] == ".docx"

# ---------- Ejecución ----------

if operation == "1":
    # Obtener lista de proyectos de Android -->
    try:
        os.chdir(ruta_de_proyectos)
    except:
        print("Esa ruta no existe. Verifica el nombre de usuario.")
        sys.exit()
    proyectos = os.listdir()

    # Cambiar a folder nuevo en escritorio -->
    try:
        os.chdir(os.path.join(escritorio, "Android a Word"))
    except:
        print("Esa ruta no existe. Verifica el nombre de usuario.")
        sys.exit()

    # Crear documentos por cada proyecto -->
    for proyecto_i in proyectos:
        document = docx.Document()

        ruta_i = os.path.join(ruta_de_proyectos, proyecto_i)
        archivos_1 = {}
        for root, dirs, files in os.walk(ruta_i):
            if filtro_por_carpeta(root, proyecto_i):
                for name in files:
                    if filtro_por_extension(name) and filtro_por_archivo(name):
                        archivos_1[name] = os.path.join(root, name)

        # Copiar contenido de archivos al documento -->
        for i in archivos_1:
            document.add_heading(i, level=1)
            if es_imagen(i):
                document.add_picture(archivos_1[i], width=docx.shared.Inches(1.25))
            elif es_mp3(i):
                document.add_paragraph("Archivo de sonido.")
            elif es_mp4(i):
                document.add_paragraph("Archivo de video.")
            else:
                f = open(archivos_1[i], "r")
                document.add_paragraph(f.read())

        document.save(proyecto_i + '.docx')
elif operation == "2":
    print("creando ruta, buscando archivos word...")
    #ruta a los documentos word
    try:
        os.chdir(os.path.join(escritorio, "Android a Word"))
    except:
        print("Esa ruta no existe")
        sys.exit()
    #lista los word dentro de la ruta
    wordList = os.listdir()
    #crea carpeta para guardar los pdf y se mueve para alla
    try:
        os.mkdir(os.path.join(escritorio, "Word to Pdf"))
    except:
        pass
    #se cambia a la ruta de los pdf
    try:
        os.chdir(os.path.join(escritorio, "Word to Pdf"))
    except:
        print("ruta para pdf no encontrada")
        sys.exit()

    #itera cada word encontrado y convierte a pdf
    for word_i in wordList:
        print(word_i)
        if es_docx(word_i):
            filename, f_ext = os.path.splitext(word_i)#recupera nombre del archivo
            docx2pdf.convert(os.path.join(escritorio, f"Android a Word//{word_i}"), os.path.join(escritorio, f"Word to Pdf//{filename}.pdf"))#nombre del pdf
